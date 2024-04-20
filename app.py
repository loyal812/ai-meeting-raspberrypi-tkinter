import pyaudio
import wave
import os
import datetime
import subprocess
import shutil
from openai import OpenAI
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import threading
import platform






meetings_folder = "meetings"

# Global variabler
global stream, frames, audio, recording, is_paused, content_length
stream = None
frames = []
audio = None
recording = False
is_paused = False  # Ny kontrolvariabel til pauser


def set_content_length(data):
    global content_length
    content_length = data
    print(content_length)

def record_loop():
    global stream, frames, recording, is_paused
    print("Optager LOOP")
    while recording:
        if not is_paused:
            try:
                data = stream.read(1024, exception_on_overflow=False)
                frames.append(data)
            except OSError as e:
                print(f"Fejl ved lydoptagelse: {e}")
                # Håndter fejlen her (fx log, genstart strøm osv.)

def toggle_pause():
    global is_paused, stream
    is_paused = not is_paused

    # Luk og genåbn audio stream ved pause/genoptag
    if is_paused:
        stream.stop_stream()
        stream.close()
    else:
        audio = pyaudio.PyAudio()
        stream = audio.open(format=pyaudio.paInt16, channels=1, rate=44100, input=True, frames_per_buffer=1024)


def start_recording():
    global stream, frames, audio, recording, is_paused
    print("Starter nu")
    frames = []
    is_paused = False

    audio = pyaudio.PyAudio()
    stream = audio.open(format=pyaudio.paInt16, channels=1, rate=44100, input=True, frames_per_buffer=1024)

    recording = True
    recording_thread = threading.Thread(target=record_loop)
    recording_thread.start()

def stop_recording(audio_folder):
    if not os.path.exists(audio_folder):
        os.makedirs(audio_folder)
        print(f"{audio_folder} has been created.")
    else:
        print(f"{audio_folder} already exists.")

    global stream, frames, audio, recording, file_path, date_str
    print("Stopper nu")

    # Stopper optagelsesloopet
    recording = False

    # Lukker og afslutter audio stream
    stream.stop_stream()
    stream.close()
    audio.terminate()

    # Behandler og gemmer den optagede lydfil
    date_str = datetime.datetime.now().strftime('%d-%m-%Y-%H-%M')
    file_path = os.path.join(audio_folder, date_str + ".wav")

    with wave.open(file_path, "wb") as audio_file:
        audio_file.setnchannels(1)
        audio_file.setsampwidth(audio.get_sample_size(pyaudio.paInt16))
        audio_file.setframerate(44100)
        audio_file.writeframes(b''.join(frames))

    print("Audio file created successfully")

    date_dir, new_file_path = date_dir_create(meetings_folder)
    mp3_file_path = convert_audio(new_file_path, date_dir)
    output_dir = split_audio(date_dir, mp3_file_path)
    transcribe_audio_files(output_dir, date_dir)
    generate_meeting_report(date_dir)
    # send_email(date_dir)
    
    return file_path




def date_dir_create(meetings_folder):
    date_str = datetime.datetime.now().strftime('%d-%m-%Y-%H-%M')
    #Making a date folder and copying the audio file into the folder
    date_dir = os.path.join(meetings_folder, date_str)
    os.makedirs(date_dir, exist_ok=True)
    print("Made a mew folder in the meetings folder called: " + date_dir)
    new_file_path = shutil.copy2(file_path, date_dir)
    return date_dir, new_file_path



def convert_audio(new_file_path, date_dir):
    # Rename the audio file in the meetings folder to Part_full.wav
    new_file_name = os.path.join(date_dir, "Part_Full.wav")
    os.rename(new_file_path, new_file_name)

    # Convert the WAV to MP3
    mp3_file_path = os.path.splitext(new_file_name)[0] + ".mp3"
    command = ["ffmpeg", "-i", new_file_name, "-vn", "-acodec", "libmp3lame", "-q:a", "2", mp3_file_path]
    subprocess.run(command, check=True)

    #makes the new Dir for the mp3 file and deletes the old audio file
    os.remove(new_file_name)
    print("Wav file converted into MP3 file successfully")
    
    return mp3_file_path


def split_audio(date_dir, mp3_file_path):
    #makes a dir called audio_files
    output_dir = os.path.join(date_dir, "audio_files")
    os.makedirs(output_dir)
    print("Output_dir called audio_files made")


    #Splitting audio

   # Set segment time to 5 minutes in seconds to aim for <10MB segments
    segment_time_seconds = 60 * 10
    output_path = os.path.join(output_dir + "/audio_part_%03d" + ".mp3")

    command = ["ffmpeg", "-i", mp3_file_path, "-f", "segment", "-segment_time", str(segment_time_seconds), "-c", "copy", output_path]
    subprocess.run(command, check=True)
        
    # Print the size of each segment
    for segment_file in os.listdir(output_dir):
        segment_path = os.path.join(output_dir, segment_file)
        segment_size = os.path.getsize(segment_path)
        print(f"Segment {segment_file}: {segment_size}")
    
    return output_dir
   



def transcribe_audio_files(output_dir, date_dir):
    client = OpenAI(api_key="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")
    
    text_output = []  # List to store transcriptions
    audio_files = [f for f in os.listdir(output_dir)]
    total_files = len(audio_files)

    # Iterate through each audio file in the directory
    for index, audio_file in enumerate(audio_files, start=1):
        audio_path = os.path.join(output_dir, audio_file)
        print(f"Transcribing file {index} of {total_files}: {audio_file}...")

        with open(audio_path, "rb") as f:
            transcript_response = client.audio.transcriptions.create(model="whisper-1", file=f)

            # Access the text attribute directly
            transcript_text = transcript_response.text

        text_output.append(transcript_text)

    # Save the combined transcriptions to a file
    with open(os.path.join(date_dir, 'transcription.txt'), 'w', encoding='utf-8') as f:
        f.write("\n".join(text_output))

    print("Transcription completed!")









def split_transcript_into_parts(date_dir):
    with open(os.path.join(date_dir, "transcription.txt"), "r") as file:
        transcript_text = file.read()

    total_length = len(transcript_text)
    part_length = total_length // 3

    parts = [
        transcript_text[:part_length], 
        transcript_text[part_length:2*part_length], 
        transcript_text[2*part_length:]
    ]
    print("Splitting of txt file done")

    return parts


def generate_meeting_report(date_dir):
    client = OpenAI(api_key="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")
    
    # client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key="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",)
   



    # Split the transcript into three parts
    transcript_parts = split_transcript_into_parts(date_dir)



    messages = [{
        "role": "system",
        "content": "You are about to receive parts of a transcribed meeting. Please process them sequentially."
    }]

    # Provide transcript parts to GPT-4
    for idx, part in enumerate(transcript_parts):
        print(f"Sending TXT data to API: {idx + 1}...")
        messages.append({
            "role": "user",
            "content": part
        })

        # Initialize conversation with GPT-4
    # response = client.chat.completions.create(model="openai/gpt-4-32k", # Optional (user controls the default)
    # messages=messages,
    # extra_headers={
    # "HTTP-Referer": "http://localhost:3000", # To identify your app. Can be set to e.g. http://localhost:3000 for testing
    # "X-Title": "AI_Meeting", # Optional. Shows on openrouter.ai
    #  })


    # Now, command GPT-4 to provide a headline
    global content_length

    if content_length == "Short":
        word_length = 250
    elif content_length == "Medium":
        word_length = 500
    elif content_length == "Long":
        word_length = 750

    messages.append(
        {
            "role": "system",
            "content": f"You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a usefull information that someone could read to quickly understand what was talked about. Is important that you at all time use the same langauge as transcripted. I need you to first of all make a headline thats is fitting for the meeting. Then i want you to make a very short summart in bullet point format. Then i want you to write a more detailed summary about the whole meeting where you focus on including key points and you are very good at gathering conclusions, dates and numbers I want you to make it easy to read with good spaces and headlines in the summary. This response should max be {word_length} words long. Its very important that you check what language the transcript is in and use the same language in your summary. The most common languange is Danish and english."
        }
    
    )

    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages
    )

    

    # response = client.chat.completions.create(model="openai/gpt-4-32k", messages=messages, extra_headers={
    #     "HTTP-Referer": "http://localhost:3000", # To identify your app. Can be set to e.g. http://localhost:3000 for testing
    #     "X-Title": "AI_Meeting", # Optional. Shows on openrouter.ai
    #  })
    Summary = completion.choices[0].message.content
    print("Done with summery. Saving to Docx file")

    save_to_docx(Summary, date_dir)


def save_to_docx(Summary, date_dir):
    doc = Document()
    doc.add_heading('Møde referat', level=0)
    doc.add_paragraph(Summary)

    # Save the doc in date_dir
    doc.save(os.path.join(date_dir, "meeting_summary.docx"))
    print('File uploaded successfully')





def send_email(date_dir, r_email):
    print("mail", date_dir)
    # r_email = "kontakt@crosscompany.dk"
    s_email = "crossitlemvig@gmail.com"
    app_password = "wrgndlzevuruadye"
    subject = "Mødereferat: " + date_dir
    message = "Her er mødereferatet fra mødet der lige har været afholdt"

    # Set up MIME
    msg = MIMEMultipart()
    msg['From'] = s_email
    msg['To'] = r_email
    msg['Subject'] = subject
    msg.attach(MIMEText(message, 'plain'))
    
    # Attach the file
    docx_filename = "meeting_summary.docx"
    docx_filepath = os.path.join(f'./meetings/{date_dir}', docx_filename)
    print(docx_filepath)
    with open(docx_filepath, "rb") as attachment:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', "attachment; filename= %s" % docx_filename)
        msg.attach(part)

    # Convert the message to string
    text = msg.as_string()

    # Connect to Gmail's SMTP server
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(s_email, app_password)
    server.sendmail(s_email, r_email, text)

    print("Email has been sent to " + r_email)
    server.quit()
    

# function to establish a new connection
def createNewConnection(name, SSID, password):
    config = """<?xml version=\"1.0\"?>
        <WLANProfile xmlns="http://www.microsoft.com/networking/WLAN/profile/v1">
            <name>"""+name+"""</name>
            <SSIDConfig>
                <SSID>
                    <name>"""+SSID+"""</name>
                </SSID>
            </SSIDConfig>
            <connectionType>ESS</connectionType>
            <connectionMode>auto</connectionMode>
            <MSM>
                <security>
                    <authEncryption>
                        <authentication>WPA2PSK</authentication>
                        <encryption>AES</encryption>
                        <useOneX>false</useOneX>
                    </authEncryption>
                    <sharedKey>
                        <keyType>passPhrase</keyType>
                        <protected>false</protected>
                        <keyMaterial>"""+password+"""</keyMaterial>
                    </sharedKey>
                </security>
            </MSM>
        </WLANProfile>"""
    command = "netsh wlan add profile filename=\""+name+".xml\""+" interface=Wi-Fi"
    with open(name+".xml", 'w') as file:
        file.write(config)
    os.system(command)
 
# function to connect to a network    
def connect(name, SSID):
    command = "netsh wlan connect name=\""+name+"\" ssid=\""+SSID+"\" interface=Wi-Fi"
    os.system(command)

# function to display avavilabe Wifi networks    
def displayAvailableNetworks():
    command = "netsh wlan show networks interface=Wi-Fi"
    os.system(command)

def get_wifi_status():
    os_type = platform.system()
    if os_type == 'Linux':
        result = subprocess.run(['iwgetid'], capture_output=True, text=True)
    elif os_type == 'Windows':
        result = subprocess.run(['netsh', 'wlan', 'show', 'interfaces'], capture_output=True, text=True)
    else:
        return "Unsupported operating system"
    # Check the return code of the subprocess
    return result.returncode == 0


def get_meeting_list():
    directory_path = './meetings'

    try:
        directory_contents = os.listdir(directory_path)
        return directory_contents
    except FileNotFoundError:
        print(f"The directory at path '{directory_path}' does not exist or could not be found.")
        return []  # Return an empty list to signify that no contents were found


def delete_meeting_data(folder_name):
    # Specify the path of the folder you want to delete
    audio_path = f'./Audio_Recordings/{folder_name}.wav'
    delete_file(audio_path)

    meeting_path = f'./meetings/{folder_name}'
    meeting_audio_path = f'./meetings/{folder_name}/audio_files'

    meeting_audio_file_list = os.listdir(meeting_audio_path)
    for file_name in meeting_audio_file_list:
        delete_file(f'./meetings/{folder_name}/audio_files/{file_name}')

    delete_folder(meeting_audio_path)

    meeting_file_list = os.listdir(meeting_path)

    for file_name in meeting_file_list:
        delete_file(f'./meetings/{folder_name}/{file_name}')

    delete_folder(meeting_path)
    

def delete_folder(folder_path):
    # Check if the folder exists before attempting to delete it
    if os.path.exists(folder_path):
        try:
            os.rmdir(folder_path)
            print(f"The folder at {folder_path} has been deleted.")
        except OSError as e:
            print(f"Error: {folder_path} : {e.strerror}")
    else:
        print(f"The folder at {folder_path} does not exist.")

def delete_file(file_path):
    # Check if the file exists before attempting to remove it
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"The file at {file_path} has been removed.")
    else:
        print(f"The file at {file_path} does not exist.")



